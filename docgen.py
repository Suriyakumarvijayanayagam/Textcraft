from docx import Document
from docx.shared import Pt, RGBColor
from docx.shared import Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION

def analyze_and_add(text):
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.page_height = Cm(30)
        section.page_width = Cm(20)
        # Create border
        border_element = section._sectPr.xpath('.//w:pgBorders')
        if border_element:
            top_border = OxmlElement('w:top')
            top_border.set(qn('w:val'), 'single')
            top_border.set(qn('w:sz'), '4')
            top_border.set(qn('w:space'), '0')
            top_border.set(qn('w:color'), 'auto')
            border_element[0].append(top_border)
            left_border = OxmlElement('w:left')
            left_border.set(qn('w:val'), 'single')
            left_border.set(qn('w:sz'), '4')
            left_border.set(qn('w:space'), '0')
            left_border.set(qn('w:color'), 'auto')
            border_element[0].append(left_border)

    lines = text.split('\n')

    for line in lines:
        if line.strip():  # Check if the line is not empty
            if line.startswith("#"):
                # If the line starts with '#', assume it's a heading
                heading_level = min(line.count('#') + 1, 6)  # Cap heading level at 6
                heading = doc.add_heading(level=heading_level)
                heading_run = heading.add_run(line.strip("#").strip().upper())   # Capitalize all words
                heading_run.font.size = Pt(20)
                heading_run.font.name = 'Skfinal'
                heading_run.font.color.rgb = RGBColor(0, 0, 0)
            elif line.startswith("*"):
                # If the line starts with '*', assume it's a heading
                heading_level = min(line.count('*') + 1, 6)  # Cap heading level at 6
                heading = doc.add_heading(level=heading_level)
                heading_run = heading.add_run(line.strip("#").strip().upper())  # Capitalize all letters
                heading_run.font.size = Pt(20)
                heading_run.font.name = 'Skfinal'
                heading_run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                # Otherwise, treat it as a regular paragraph
                para = doc.add_paragraph()
                para_run = para.add_run(line)
                para_run.font.size = Pt(20)
                para_run.font.name = 'Skfinal'
                para_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    return doc
